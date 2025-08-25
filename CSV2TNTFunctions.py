#Bailey Mark Weiss 26 January 2025
#   updated, cleaned, and annotated 24 August 2025

#Python functions to translate google sheets to TNT file for matrices and Docx file for character statements
#Read Manual for explanation on layout of sheets
#Ensure that the Google Sheet is shared "Anyone with the link can edit"

#-----------------------Import packages-----------------------
import csv
import requests
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.shared import Cm  # For setting margins in centimeters
from docx.oxml.ns import qn  # For XML namespace
from docx.oxml import OxmlElement  # For adding columns
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
import easygui
#---------------------------------------------------


#-----------------------FUNCTION - PrevInfo() - reads and returns information from InfoFile-----------------------
def PrevInfo(path): #accepts a path as a string to the folder the file is in
    UsePrev = None
    if not os.path.exists(path+"/infoFile.txt"): #if no file exhists
        print("No previous Info File")
        UsePrev = "N"
        return UsePrev, "None", "None", "None", "None" #all set to none
    else:
        PrevInfo = []
        with open(path+"/InfoFile.txt") as csvfile: #reads in file information
            csvdata = csv.reader(csvfile, delimiter='#')
            for row in csvdata:
                PrevInfo.append(row)
        while len(PrevInfo) < 4:
            PrevInfo.append(["None"] * 2)
        print("InfoFile read")
        if PrevInfo[0][1] == "None" or PrevInfo[1][1] == "None": #checks if Name or Google Sheet Key is empty
            print("Previous file is empty!")
            UsePrev = "N"
        return UsePrev, PrevInfo[0][1], PrevInfo[1][1], PrevInfo[2][1], PrevInfo[3][1] #returns data as a tupple - UsePrev, Name, Google Sheets Key, Statements GID, Characters GID
#---------------------------------------------------


#-----------------------FUNCTION - MakeFolder() - makes a folder at the directory-----------------------     
def MakeFolder(path): #accepts a path as a string to the directory the folder will be created in
    if not os.path.exists(path):
        os.mkdir(path) #creates folder for files if not already there
        print("Folder %s created!" % path)
    else:
        print("Folder %s already exists" % path)
#---------------------------------------------------


#-----------------------FUNCTION - load_data() - translates a CSV to a dataset-----------------------
def load_data(filename): #accepts a string with the name of the file
    dataset = []
    with open(filename, encoding="utf8") as csvfile:
        csvdata = csv.reader(csvfile, delimiter=',')
        for row in csvdata:
            dataset.append(row)
        return dataset
#---------------------------------------------------


#-----------------------FUNCTION - downloadCSV() - Downloads a csv file from Google Sheets-----------------------
def downloadCSV(GSKey, GID, CSV): #accepts 3 strings: Google Sheets Key, the GID of the sheet, and the name of the file to save it to
    url = "https://docs.google.com/spreadsheets/d/"+GSKey+"/export?format=csv&gid="+GID #creates the URL to access
    try:
        response = requests.get(url, stream=True)
        response.raise_for_status() #checks if the request was successful
        with open(CSV, "wb")as file:
            file.write(response.content) #writes data to csv file
        print(f"Characters downloaded successfully")
    except requests.exceptions.RequestException as e:
        print(f"An error occurred: {e}")
#---------------------------------------------------


#-----------------------FUNCTION - numRows() - gets the number of rows in a dataset-----------------------
def numRows(dataset): #accepts a dataset
    return len(dataset) #returns the number of rows
#---------------------------------------------------


#-----------------------FUNCTION - numColumns() - gets the number of columns in a dataset-----------------------
def numColumns(dataset): #accepts a dataset
    return len(dataset[0]) #returns the number of columns
#---------------------------------------------------


#-----------------------FUNCTION - createTNT()- creates a tnt file of the characters-----------------------
def createTNT(path, TNTFileName, GSKey, GID): #accepts 4 strings: directory the file will be saved in, name of the TNT file, Google Sheets key, GID of the sheet
    CharacterCSV = path+"/"+TNTFileName+"_Characters.csv" #creates full path for the CSV file
    RemoveChar = ("/","&") #These characters will be removed and [] will be added around the score

    downloadCSV(GSKey,GID,CharacterCSV) #downloads the characters file
    matrix = load_data(CharacterCSV) #loads csv into a dataset

    ordered = matrix[0] #first row denotes which characters are ordered, saved into seperate dataset
    del matrix[0] #delete the first row
    numCharacters = numColumns(matrix)-1 #calculates number of characters
    numTaxa = numRows(matrix)-1 #calculates number of taxa

    #------------------Loops through matrix and replaces characters------------------
    for i in range(len(matrix)):
        for j in range(len(matrix[i])):
            if matrix[i][j] is None or matrix[i][j] == "":
                matrix[i][j] ="?" #replaces empty cells in the matrix with ?
            for char in RemoveChar:
                if char in matrix[i][j]:
                    matrix[i][j] = '['+matrix[i][j].replace(char, ' ')+']' #replaces predefined characters with a space and adds [] around the score
    #------------------------------------
    
    #------------------Loops through taxon names and removes spaces------------------
    for i in range(len(matrix)):
        if " " in matrix[i][0]:
                matrix[i][0] = matrix[i][0].replace(" ", "_") #replaces spaces in taxon names with underscores
    #------------------------------------

    #------------------Creates the preamble of the TNT file------------------
    text = (
        "xread\n'TNT file created from Google Sheets on "+str(datetime.today().strftime('%d %B %Y'))
        +" using:\n         CSV2TNT\n           by\n     Bailey Weiss, 2025'\n"
        +str(numCharacters)+' '+str(numTaxa)+'\n') #writes the number of characters and taxa in the file
    #------------------------------------

    #------------------Loops through matrix and writes taxa and characters in the correct format------------------
    for i in range(1, len(matrix)):
        list = matrix[i][0] +'\t' + ''.join(map(str, matrix[i][1:numCharacters+1])) + '\n'
        text = text + list
    #------------------------------------

    #------------------Loops through character names and writes them in the corect format------------------
    text=text+";\n\ncnames\n" #syntax
    for i in range(1, len(matrix[0])):
        text =text +'{'+str(i-1)+' Character'+ matrix[0][i] + ';\n'
    #------------------------------------
    
    #------------------Loops through ordered list and writes them in the corect format------------------
    text=text + ';\n\nccode\t +  ' #syntax
    del ordered[0] #deletes row name
    for i in range(len(ordered)):
        if ordered[i] == "*": #orded characters are denoted with a * in the csv
            text=text+str(i) + ' '
    #------------------------------------

    #------------------Writes the text to the TNT file------------------
    text=text+'*;\n\n\nproc /;\ncomments 0\n;' #syntax
    TNTfile = open(path+"/"+TNTFileName+"_Matrix_"+str(datetime.today().strftime('%d_%b_%y'))+".tnt", 'w')
    TNTfile.write(text)
    TNTfile.close()
    #------------------------------------
    
    print("TNT file created")
#---------------------------------------------------


#-----------------------FUNCTION - CharStatements()- creates a docx file of the statments-----------------------
def CreateCharStatements(path, DocFileName, GSKey, GID): #accepts 4 strings: directory the file will be saved in, name of the docx file, Google Sheets key, GID of the sheet
    StatementCSV = path+"/"+"Statements.csv" #creates full path for the CSV file
    
    downloadCSV(GSKey,GID,StatementCSV) #downloads the statements file
    statements = load_data(StatementCSV) #loads csv into a dataset

    #------------------Sets format of docx file------------------
    doc = Document()
    style = doc.styles['Normal'] #document style
    font = style.font
    font.name = 'Times New Roman'  #font name
    font.size = Pt(12)  #font size
    #------------------------------------
    
    #------------------Gets the column number of the elements and statments------------------
    StatementsColumnText = "" #initialise variable
    StatementsColumnNumber = 0 #initialise variable

    for i in range(len(statements[0])):
        StatementsColumnText = StatementsColumnText + str(i) + ": "+ statements[0][i] + "\n" #list of column names
    ElementsColumnNumber = str(easygui.enterbox("Type the number of the column containing the element names:\n"+ StatementsColumnText)) #selects the column containing the element names, will be used to group statements
    StatementsColumnNumber = str(easygui.enterbox("Type the number of the column containing the statements:\n"+ StatementsColumnText)) #selects the colmun containing the statments
    #------------------------------------
    
    #------------------Loops through statements and formates them correctly and creates headings based on elements------------------
    Element = "" #initialise variable
    for i in range(1, len(statements)): #checks each rows, element column and then statement column
        if statements[i][int(ElementsColumnNumber)]!=Element and statements[i][0]!="999" and statements[i][0]!="": #creates a heading if the previous element is different to the current element
            if Element!="":
                doc.add_paragraph()
            Element=statements[i][int(ElementsColumnNumber)] #name of section i.e. element name
            heading = doc.add_heading(Element, level=1)
            run = heading.runs[0] #formating
            run.font.name = 'Times New Roman' #formating
            run.font.size = Pt(12) #formating
            run.font.color.rgb = RGBColor(0, 0, 0) #formating
        if statements[i][0]=="999" and Element != "Removed characters": #creates a removed characters paragraph if they exist
            Element = "Removed characters" #name of section
            heading = doc.add_heading(Element, level=1)
            run = heading.runs[0] #formating
            run.font.name = 'Times New Roman' #formating
            run.font.size = Pt(12) #formating
            run.font.color.rgb = RGBColor(0, 0, 0) #formating
        doc.add_paragraph(statements[i][int(StatementsColumnNumber)]) #writes the statement
    #------------------------------------

    #------------------writes and formats the docx file------------------
    for paragraph in doc.paragraphs:
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Pt(0) #removes the space before the paragraph
        paragraph_format.space_after = Pt(0) #removes the space after the paragraph
    section = doc.sections[0]
    section.left_margin = Cm(1.3) #left margin
    section.right_margin = Cm(1.3) #right margin
    section.top_margin = Cm(1.3) #top margin
    section.bottom_margin = Cm(1.3) #bottom margin
    sectPr = section._sectPr #access to the section properties
    cols = sectPr.xpath('./w:cols')[0] if sectPr.xpath('./w:cols') else OxmlElement('w:cols') #check if <w:cols> exists
    cols.set(qn('w:num'), '2') #sets the number of columns to 2
    sectPr.append(cols) #adds the column element back to the section properties
    for paragraph in doc.paragraphs:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY #justifies the document
    doc.save(path+"/"+DocFileName+"_Statements_"+str(datetime.today().strftime('%d_%b_%y'))+".docx") #saves the file
    os.remove(StatementCSV)
    #------------------------------------

    print("DOCX file created")
#---------------------------------------------------